import re
from docx import Document
import aspose.words as aw
import parcer

# document = Document("Output.docx")



# doc = aw.Document("Input.doc")
# doc.save("Output.docx")


# Регулярка для поиска последовательностей пробелов: от двух подряд и более


def get_table(group):

    parcer.download_docx(url ='https://vgpgk.ru/raspisanie/vgpgk-zameny-1-korpus.doc')
    document = Document("raspisanie.docx")

    # multi_space_pattern = re.compile(r'\s{2,}')

    for table in document.tables:
        for x in range(0, 2):
            row = table.column_cells(x)
            gg = [re.compile(r'\s{2,}').sub(' ', i.text.strip()) for i in row]
            for i in range(0, len(gg), 2):
                if gg[i] == group:
                    return(f'Расписание для группы на {gg[i]}❗️\n{str(gg[i+1])}\n')
    return False