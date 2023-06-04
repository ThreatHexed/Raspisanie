from calendar import weekday
import re
from docx import Document
import aspose.words as aw
import parcer
import newraspisanie



def get_table(group):
    global data
    parcer.download_docx(url ='https://vgpgk.ru/raspisanie/vgpgk-zameny-1-korpus.doc')
    document = Document("raspisanie.docx")
    aboba = []

    for paragraph in document.paragraphs:
        
        if "на " in paragraph.text:
            aboba.append(paragraph.text)
        
    for table in enumerate(document.tables, start=0):
        for x in range(0, 3):
            row = table[1].column_cells(x)
            
            gpz = [re.compile(r'\s{2,}').sub(' ', i.text.strip()) for i in row]
            for i in range(0, len(gpz), 2):
                if gpz[i] == group:
                    if table[0] == 0 or table[0] == 1:
                        
                        data = aboba[0]
                        for table2 in enumerate(document.tables[2:-1], start=0):
                            for y in range(0, 3):
                                row2 = table2[1].column_cells(y)
                                gpz2 = [re.compile(r'\s{2,}').sub(' ', y.text.strip()) for y in row2]

                                for t in range(0, len(gpz2), 2):
                                    if gpz2[t] == group:
                                        return(f'Расписание для группы {gpz[i]} {data}❗️\n{str(gpz[i+1])}\n\nРасписание для группы {aboba[-1]}:\n{gpz2[t+1]}')
                                        
                        return newraspisanie.pars(data, gpz[i+1])

                    elif table[0] == 2 or table[0] == 3:
                        data = aboba[2]
                        return(f'Расписание для группы {gpz[i]} {data}❗️\n{str(gpz[i+1])}')
                    

    

    return newraspisanie.pars(data, 0)


get_table('ИС-213')